"""
Converte LEVANTAMENTO_METADADOS.md para LEVANTAMENTO_METADADOS.docx
com formatação Word completa (títulos, tabelas, negrito, itálico, blockquotes).

Uso: python3 scripts/md_to_docx.py
"""

import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Caminhos ────────────────────────────────────────────────────────────────
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC  = os.path.join(BASE_DIR, "LEVANTAMENTO_METADADOS.md")
DEST = os.path.join(BASE_DIR, "LEVANTAMENTO_METADADOS.docx")

# ─── Paleta de cores ─────────────────────────────────────────────────────────
COR_TITULO_DOC  = RGBColor(0x1a, 0x3a, 0x5c)   # azul escuro
COR_H1          = RGBColor(0x1a, 0x3a, 0x5c)
COR_H2          = RGBColor(0x1e, 0x6a, 0x9e)
COR_H3          = RGBColor(0x2e, 0x86, 0xc1)
COR_BLOCKQUOTE  = RGBColor(0x55, 0x55, 0x55)
COR_HEADER_BG   = RGBColor(0x1e, 0x6a, 0x9e)   # fundo do cabeçalho de tabela
COR_ROW_ALT     = RGBColor(0xf0, 0xf6, 0xfc)   # azul bem claro
COR_SEPARATOR   = RGBColor(0xbb, 0xdd, 0xf0)


def set_cell_bg(cell, color: RGBColor):
    """Define cor de fundo de uma célula."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """Define bordas individuais de uma célula. kwargs: top, bottom, left, right."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, specs in kwargs.items():
        tag = OxmlElement(f"w:{edge}")
        for attr, val in specs.items():
            tag.set(qn(f"w:{attr}"), val)
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_run_with_inline(para, text: str, base_bold=False, base_italic=False,
                        base_color=None, base_size=None):
    """
    Adiciona runs a um parágrafo interpretando **negrito**, *itálico* e `code`.
    """
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    pos = 0
    for m in pattern.finditer(text):
        # Texto antes do match
        before = text[pos:m.start()]
        if before:
            r = para.add_run(before)
            r.bold   = base_bold
            r.italic = base_italic
            if base_color: r.font.color.rgb = base_color
            if base_size:  r.font.size = base_size

        full = m.group(0)
        if full.startswith("**"):
            r = para.add_run(m.group(2))
            r.bold   = True
            r.italic = base_italic
        elif full.startswith("*"):
            r = para.add_run(m.group(3))
            r.bold   = base_bold
            r.italic = True
        else:  # backtick code
            r = para.add_run(m.group(4))
            r.font.name = "Courier New"
            r.font.size = Pt(9)
            r.bold = False

        if base_color: r.font.color.rgb = base_color
        if base_size and not full.startswith("`"):
            r.font.size = base_size
        pos = m.end()

    # Texto restante
    tail = text[pos:]
    if tail:
        r = para.add_run(tail)
        r.bold   = base_bold
        r.italic = base_italic
        if base_color: r.font.color.rgb = base_color
        if base_size:  r.font.size = base_size


def style_paragraph(para, size=None, color=None, bold=False, italic=False,
                    space_before=0, space_after=6, alignment=None):
    """Aplica formatação ao parágrafo inteiro."""
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if alignment:
        para.alignment = alignment
    for run in para.runs:
        if size:  run.font.size  = Pt(size)
        if color: run.font.color.rgb = color
        run.bold   = bold
        run.italic = italic


def build_table(doc, headers, rows):
    """Cria tabela formatada com cabeçalho colorido e linhas zebradas."""
    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Cabeçalho
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, COR_HEADER_BG)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        # strip inline markdown from header
        clean_h = re.sub(r'[*`]', '', h)
        r = p.add_run(clean_h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9)

    # Linhas de dados
    for row_idx, row_data in enumerate(rows):
        tr = table.rows[row_idx + 1]
        bg = COR_ROW_ALT if row_idx % 2 == 1 else None
        for col_idx, cell_text in enumerate(row_data):
            cell = tr.cells[col_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if bg:
                set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            add_run_with_inline(p, cell_text, base_size=Pt(9))

    # Largura automática
    for col in table.columns:
        for cell in col.cells:
            cell.width = Inches(6.5 / n_cols)

    doc.add_paragraph()  # espaço após tabela


def parse_and_build(doc, lines):
    """Percorre as linhas do Markdown e constrói o documento Word."""
    i = 0
    in_code_block = False
    code_lines    = []

    while i < len(lines):
        line = lines[i]

        # ── Bloco de código (``` ... ```) ─────────────────────────────────────
        if line.strip().startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                # Renderiza como parágrafo monoespaçado com fundo cinza
                for cl in code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent  = Cm(0.5)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after  = Pt(0)
                    r = p.add_run(cl)
                    r.font.name = "Courier New"
                    r.font.size = Pt(8.5)
                    r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
                doc.add_paragraph()
            i += 1
            continue

        if in_code_block:
            code_lines.append(line.rstrip())
            i += 1
            continue

        # ── Separadores horizontais (---) ─────────────────────────────────────
        if re.match(r'^-{3,}\s*$', line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"),
                       f"{COR_SEPARATOR[0]:02X}{COR_SEPARATOR[1]:02X}{COR_SEPARATOR[2]:02X}")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # ── Tabelas Markdown (| col | col |) ─────────────────────────────────
        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1

            # Separa header da linha de separação (---|---) e das linhas de dados
            if len(table_lines) < 2:
                continue

            # Parse do header
            raw_headers = [c.strip() for c in table_lines[0].split("|") if c.strip()]

            # Linha 1 é o separador (---|---), pula
            data_rows = []
            for tl in table_lines[2:]:
                cells = [c.strip() for c in tl.split("|") if c.strip() != ""]
                # Garante mesmo número de colunas
                while len(cells) < len(raw_headers):
                    cells.append("")
                data_rows.append(cells[:len(raw_headers)])

            build_table(doc, raw_headers, data_rows)
            continue

        # ── Títulos ───────────────────────────────────────────────────────────
        h_match = re.match(r'^(#{1,4})\s+(.*)', line)
        if h_match:
            level   = len(h_match.group(1))
            content = h_match.group(2).strip()

            # Título principal do documento (# único)
            if level == 1:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(10)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                add_run_with_inline(p, content, base_bold=True,
                                    base_color=COR_H1, base_size=Pt(20))
            elif level == 2:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after  = Pt(4)
                add_run_with_inline(p, content, base_bold=True,
                                    base_color=COR_H2, base_size=Pt(14))
                # Linha sob o H2
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "4")
                bottom.set(qn("w:space"), "1")
                bottom.set(qn("w:color"),
                           f"{COR_H2[0]:02X}{COR_H2[1]:02X}{COR_H2[2]:02X}")
                pBdr.append(bottom)
                pPr.append(pBdr)
            elif level == 3:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after  = Pt(3)
                add_run_with_inline(p, content, base_bold=True,
                                    base_color=COR_H3, base_size=Pt(12))
            else:  # H4
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after  = Pt(2)
                add_run_with_inline(p, content, base_bold=True,
                                    base_color=COR_H3, base_size=Pt(11))
            i += 1
            continue

        # ── Blockquote (> ...) ────────────────────────────────────────────────
        if line.strip().startswith(">"):
            content = re.sub(r'^>\s?', '', line.strip())
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Cm(1.0)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            # Borda esquerda colorida
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            left_bdr = OxmlElement("w:left")
            left_bdr.set(qn("w:val"), "single")
            left_bdr.set(qn("w:sz"), "12")
            left_bdr.set(qn("w:space"), "4")
            left_bdr.set(qn("w:color"),
                         f"{COR_H2[0]:02X}{COR_H2[1]:02X}{COR_H2[2]:02X}")
            pBdr.append(left_bdr)
            pPr.append(pBdr)

            if content == "[ PREENCHER ]":
                r = p.add_run("[ PREENCHER ]")
                r.font.color.rgb = RGBColor(0xCC, 0x44, 0x00)
                r.bold   = True
                r.italic = True
                r.font.size = Pt(10)
            else:
                add_run_with_inline(p, content, base_italic=True,
                                    base_color=COR_BLOCKQUOTE, base_size=Pt(10))
            i += 1
            continue

        # ── Lista não-ordenada (- item) ───────────────────────────────────────
        if re.match(r'^[-*]\s+', line.strip()):
            content = re.sub(r'^[-*]\s+', '', line.strip())
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            add_run_with_inline(p, content, base_size=Pt(10))
            i += 1
            continue

        # ── Linha em branco ───────────────────────────────────────────────────
        if not line.strip():
            # Pequeno espaço sem criar parágrafos extras desnecessários
            i += 1
            continue

        # ── Metadados do cabeçalho (**Campo:** valor) ─────────────────────────
        meta_match = re.match(r'^\*\*(.+?):\*\*\s*(.*)', line.strip())
        if meta_match:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            label = meta_match.group(1)
            value = meta_match.group(2)
            r = p.add_run(f"{label}: ")
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = COR_H1
            add_run_with_inline(p, value, base_size=Pt(10))
            i += 1
            continue

        # ── Parágrafo normal ──────────────────────────────────────────────────
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(4)
        add_run_with_inline(p, line.strip(), base_size=Pt(10.5))
        i += 1


def main():
    with open(SRC, encoding="utf-8") as f:
        lines = f.readlines()
    # Remove \n de cada linha mas mantém o conteúdo
    lines = [l.rstrip("\n") for l in lines]

    doc = Document()

    # ── Configuração de página ────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width   = Cm(21.0)
    section.page_height  = Cm(29.7)
    section.left_margin  = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin   = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    # ── Estilo padrão do documento ────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    parse_and_build(doc, lines)

    doc.save(DEST)
    print(f"✅  Documento gerado: {DEST}")


if __name__ == "__main__":
    main()
